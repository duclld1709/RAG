"""Document management endpoints."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from pypdf import PdfReader

from app.api.dependencies import get_document_service
from app.api.schemas.documents import DocumentDetailResponse, DocumentResponse, DocumentUpsertRequest
from app.domain.models.document import DocumentRecord
from app.services.document_service import DocumentService

router = APIRouter()

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


@router.get("/", response_model=list[DocumentResponse])
def list_documents(document_service: DocumentService = Depends(get_document_service)) -> list[DocumentResponse]:
    documents = document_service.list_documents()
    return [to_response(doc) for doc in documents]


@router.post("/", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    filename: str | None = Form(None),
    document_service: DocumentService = Depends(get_document_service),
) -> DocumentResponse:
    display_name = (filename or file.filename or "").strip()
    if not display_name:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Filename is required.")

    try:
        content = await extract_text_from_upload(file)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to process document.") from exc
    finally:
        await file.close()

    try:
        record = document_service.upsert_document(filename=display_name, content=content)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc

    return to_response(record)


@router.put("/{document_id}", response_model=DocumentResponse)
def update_document(
    document_id: UUID,
    payload: DocumentUpsertRequest,
    document_service: DocumentService = Depends(get_document_service),
) -> DocumentResponse:
    record = document_service.get_document(document_id)
    if not record:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    try:
        updated = document_service.upsert_document(
            filename=payload.filename,
            content=payload.content,
            document_id=document_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
    return to_response(updated)


@router.get("/{document_id}", response_model=DocumentDetailResponse)
def get_document(
    document_id: UUID,
    document_service: DocumentService = Depends(get_document_service),
) -> DocumentDetailResponse:
    result = document_service.get_document_with_content(document_id)
    if not result:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    record, content = result
    return to_response(record, content=content)


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    document_id: UUID,
    document_service: DocumentService = Depends(get_document_service),
) -> None:
    record = document_service.get_document(document_id)
    if not record:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    document_service.delete_document(document_id)


def to_response(
    record: DocumentRecord,
    *,
    content: str | None = None,
) -> DocumentResponse | DocumentDetailResponse:
    payload = {
        "id": record.id,
        "filename": record.filename,
        "summary": record.summary,
        "uploaded_at": record.uploaded_at,
        "updated_at": record.updated_at,
    }
    if content is not None:
        return DocumentDetailResponse(**payload, content=content)
    return DocumentResponse(**payload)


async def extract_text_from_upload(upload: UploadFile) -> str:
    extension = Path(upload.filename or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Allowed extensions: .txt, .pdf, .docx.")

    raw_bytes = await upload.read()
    if not raw_bytes:
        raise ValueError("Uploaded file is empty.")

    if extension == ".txt":
        try:
            text = raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = raw_bytes.decode("utf-8", errors="ignore")
        content = text.strip()
    elif extension == ".pdf":
        reader = PdfReader(BytesIO(raw_bytes))
        pages: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted:
                pages.append(extracted.strip())
        content = "\n\n".join(pages).strip()
    else:  # .docx
        document = DocxDocument(BytesIO(raw_bytes))
        segments: list[str] = []
        segments.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    segments.append("\t".join(cells))
        content = "\n\n".join(segments).strip()

    if not content:
        raise ValueError("No readable text content found in the uploaded file.")

    return content
